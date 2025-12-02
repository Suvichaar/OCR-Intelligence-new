#!/usr/bin/env python3
import os, sys, json, base64, requests, argparse, re, io, shutil, tempfile
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

# ==== OCR config ==============================================================
MISTRAL_OCR_ENDPOINT = os.environ.get("MISTRAL_OCR_ENDPOINT")
MISTRAL_API_KEY      = os.environ.get("API_KEY")
MISTRAL_MODEL        = os.environ.get("MISTRAL_MODEL", "mistral-document-ai-2505")

# ==== Optional Pandoc (for LaTeX math & tables) ===============================
HAVE_PYPANDOC = True
try:
    import pypandoc  # type: ignore
except Exception:
    HAVE_PYPANDOC = False

def ensure_pandoc_available() -> bool:
    if shutil.which("pandoc"):
        return True
    if HAVE_PYPANDOC:
        try:
            pypandoc.download_pandoc()
            return True
        except Exception:
            return False
    return False

# ==== Imaging / PDF / DOCX fallback ==========================================
from PIL import Image
import fitz  # PyMuPDF

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ==== Basics =================================================================
def die(msg: str, code: int = 1):
    print(f"[ERR] {msg}", file=sys.stderr); sys.exit(code)

def post_ocr(payload: dict):
    if not MISTRAL_OCR_ENDPOINT or not MISTRAL_API_KEY:
        die("Set MISTRAL_OCR_ENDPOINT and MISTRAL_API_KEY in your environment (.env).")
    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    print(f"[POST] {MISTRAL_OCR_ENDPOINT}")
    r = requests.post(MISTRAL_OCR_ENDPOINT, headers=headers, json=payload, timeout=600)
    print(f"[HTTP] {r.status_code} ctype={r.headers.get('content-type')}")
    if r.status_code >= 400:
        print("[BODY]", r.text[:4000])
        r.raise_for_status()
    return r.json()

def bytes_to_data_url(mime: str, data: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(data).decode('utf-8')}"

def _strip(s: str) -> str:
    return (s or "").strip()

def unwrap_container(resp: Dict[str, Any]) -> Dict[str, Any]:
    node = resp
    for k in ("output","response","result","data","document"):
        if isinstance(node, dict) and isinstance(node.get(k), dict):
            node = node[k]
    return node

# ==== Extract text from OCR pages ============================================
def extract_from_page(p: Dict[str, Any]) -> str:
    md = p.get("markdown")
    if isinstance(md, str) and md.strip():
        return md.strip()

    candidates: List[str] = []

    if isinstance(p.get("lines"), list):
        parts = []
        for ln in p["lines"]:
            if isinstance(ln, dict):
                t = _strip(ln.get("content") or ln.get("text"))
                if t: parts.append(t)
        if parts: candidates.append("\n".join(parts))

    if isinstance(p.get("paragraphs"), list):
        parts = []
        for para in p["paragraphs"]:
            if isinstance(para, dict):
                t = _strip(para.get("content") or para.get("text"))
                if t: parts.append(t)
        if parts: candidates.append("\n".join(parts))

    for key in ("blocks","items","elements","regions"):
        arr = p.get(key)
        if isinstance(arr, list) and arr:
            parts = []
            for it in arr:
                if isinstance(it, dict):
                    t = _strip(it.get("text") or it.get("content") or it.get("value"))
                    if t: parts.append(t)
            if parts: candidates.append("\n".join(parts))

    t = _strip(p.get("content") or p.get("text") or p.get("full_text") or p.get("raw_text"))
    if t: candidates.append(t)

    candidates = [c for c in candidates if _strip(c)]
    return max(candidates, key=len) if candidates else ""

# ==== Image I/O from OCR ======================================================
_DATA_URL_RE = re.compile(r'^data:(?P<mime>[^;]+);base64,(?P<b64>.+)$', re.IGNORECASE)

def save_image_bytes_clean(raw: bytes, out_path: Path) -> Path:
    """Normalize to PNG/JPEG; RGB; return final path with correct ext."""
    with Image.open(io.BytesIO(raw)) as im:
        fmt = (im.format or "").upper()
        if fmt == "PNG" and im.mode == "RGB":
            out = out_path.with_suffix(".png")
            im.save(out)
            return out
        im = im.convert("RGB")
        out = out_path.with_suffix(".jpg")
        im.save(out, quality=92, optimize=True)
        return out

def save_base64_image_unknown(input_str: str, out_dir: Path, stem: str) -> Path:
    """
    Accepts either a bare base64 string or a full data: URL and saves a normalized image.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    s = input_str.strip()
    m = _DATA_URL_RE.match(s)
    if m:
        raw = base64.b64decode(m.group("b64"), validate=False)  # <-- correct 'b64'
    else:
        # bare b64 OR string that still contains 'base64,'
        try:
            raw = base64.b64decode(s, validate=False)
        except Exception:
            raw = base64.b64decode(s.split("base64,", 1)[-1], validate=False)
    return save_image_bytes_clean(raw, out_dir / stem)

def fetch_and_save_http_image(url: str, out_dir: Path, stem: str) -> Optional[Path]:
    """
    If OCR returns http(s) image URLs instead of base64/data URLs, fetch & save.
    """
    try:
        r = requests.get(url, timeout=60)
        r.raise_for_status()
        raw = r.content
        try:
            with Image.open(io.BytesIO(raw)) as im:
                fmt = (im.format or "").upper()
                if fmt == "PNG" and im.mode == "RGB":
                    out = out_dir / f"{stem}.png"
                    im.save(out)
                    return out
                im = im.convert("RGB")
                out = out_dir / f"{stem}.jpg"
                im.save(out, quality=92, optimize=True)
                return out
        except Exception:
            out = out_dir / f"{stem}.bin"
            out.write_bytes(raw)
            return out
    except Exception as e:
        print(f"[WARN] fetch failed {url}: {e}")
        return None

# ==== PDF render + crops ======================================================
def render_pdf_page_to_image(pdf_path: Path, page_num: int, dpi: int = 300) -> Tuple[Image.Image, Tuple[float,float]]:
    doc = fitz.open(pdf_path.as_posix())
    try:
        page = doc[page_num-1]
        zoom = dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img  = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return img, (page.rect.width, page.rect.height)
    finally:
        doc.close()

def clamp(v, lo, hi): return max(lo, min(hi, v))

def bbox_to_pixels(b: Tuple[float,float,float,float],
                   img_w: int, img_h: int,
                   pts_wh: Optional[Tuple[float,float]] = None,
                   coord_type: str = "norm",
                   origin: str = "top-left") -> Tuple[int,int,int,int]:
    x0,y0,x1,y1 = b
    if coord_type == "norm":
        if origin == "bottom-left":
            y0, y1 = 1 - y0, 1 - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*img_w)), int(round(y0*img_h)), int(round(x1*img_w)), int(round(y1*img_h))
    elif coord_type == "pixel":
        if origin == "bottom-left":
            y0, y1 = img_h - y0, img_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0)), int(round(y0)), int(round(x1)), int(round(y1))
    elif coord_type == "pdf_points":
        if not pts_wh: raise ValueError("pdf_points requires page size in points")
        pts_w, pts_h = pts_wh; sx, sy = img_w/pts_w, img_h/pts_h
        if origin == "bottom-left":
            y0, y1 = pts_h - y0, pts_h - y1; y0, y1 = y1, y0
        X0,Y0,X1,Y1 = int(round(x0*sx)), int(round(y0*sy)), int(round(x1*sx)), int(round(y1*sy))
    else:
        raise ValueError("coord_type must be norm|pixel|pdf_points")
    x0p,x1p = sorted([X0,X1]); y0p,y1p = sorted([Y0,Y1])
    x0p = clamp(x0p,0,img_w-1); x1p = clamp(x1p,1,img_w)
    y0p = clamp(y0p,0,img_h-1); y1p = clamp(y1p,1,img_h)
    return x0p,y0p,x1p,y1p

def crop_and_save(img: Image.Image, bbox_px: Tuple[int,int,int,int],
                  padding: int, out_dir: Path, stem: str) -> Path:
    x0,y0,x1,y1 = bbox_px
    if padding:
        x0 = max(0, x0-padding); y0 = max(0, y0-padding)
        x1 = min(img.width, x1+padding); y1 = min(img.height, y1+padding)
    crop = img.crop((x0,y0,x1,y1)).convert("RGB")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{stem}.jpg"
    crop.save(out_path, quality=92, optimize=True)
    return out_path

# Always-on: if no regions.json provided, generate a full-page crop per page
def autocrops_full_pages(pdf_path: Path, dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    out: Dict[int, List[Path]] = {}
    doc = fitz.open(pdf_path.as_posix())
    try:
        for i in range(len(doc)):
            page_num = i + 1
            img, _ = render_pdf_page_to_image(pdf_path, page_num, dpi=dpi)
            saved = crop_and_save(img, (0, 0, img.width, img.height), 0, assets_dir, f"p{page_num}_full")
            out.setdefault(page_num, []).append(saved)
    finally:
        doc.close()
    print(f"[INFO] auto-crops (full pages) generated: {sum(len(v) for v in out.values())}")
    return out

# ==== Markdown helpers ========================================================
def clean_markdown(md: str) -> str:
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    return md

_MATH_OR_TABLE_RE = re.compile(
    r"(\$\$.*?\$\$|\$[^$\n]+\$|\\\(|\\\)|\\\[|\\\]|\\begin\{(equation|align|eqnarray|gather|aligned)\}|(^\s*\|.*\|\s*$\n^\s*\|?\s*[-:]+\s*(\|[-:]+\s*)+$))",
    re.MULTILINE | re.DOTALL
)

def detect_math_or_tables(pages_text: List[str]) -> bool:
    return bool(_MATH_OR_TABLE_RE.search("\n\n".join(pages_text)))

_IMG_MD_RE   = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.IGNORECASE)
_IMG_HTML_RE = re.compile(r'<img\b[^>]*>', re.IGNORECASE)

def strip_inline_images(md: str) -> str:
    return _IMG_HTML_RE.sub('', _IMG_MD_RE.sub('', md))

def clean_markdown_for_docx(md: str) -> str:
    """Clean markdown text for DOCX conversion - remove malformed image links, HTML tags, etc."""
    # Remove all image links (including malformed ones)
    md = re.sub(r'!\[[^\]]*\]\s*[^\]]*\]\([^)]*\)', '', md, flags=re.IGNORECASE)  # Malformed like ![img-1.jpeg] peg](img-1 -1.jpeg)
    md = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', md, flags=re.IGNORECASE)  # Normal image links
    
    # Convert <br> tags to line breaks
    md = md.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    md = md.replace('<br>', '\n').replace('</br>', '\n')
    
    # Remove any remaining HTML tags
    md = re.sub(r'<[^>]+>', '', md)
    
    # Clean up multiple blank lines
    md = re.sub(r'\n{3,}', '\n\n', md)
    
    return md.strip()

def md_image(path: Path, width_in: float) -> str:
    p = Path(path).resolve().as_posix()
    return f'![]({p}){{width={width_in}in}}'

def build_markdown(pages_text: List[str],
                   images_by_page: Dict[int, List[Path]],
                   crops_by_page: Dict[int, List[Path]],
                   insert_page_breaks: bool,
                   image_max_width_in: float) -> str:
    parts: List[str] = []
    for i, txt in enumerate(pages_text, start=1):
        text_clean = strip_inline_images(_strip(txt))
        parts.append(f"\n\n## Page {i}\n\n{text_clean}\n")
        for im in images_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        for im in crops_by_page.get(i, []):
            parts.append("\n" + md_image(im, image_max_width_in) + "\n")
        if insert_page_breaks and i < len(pages_text):
            parts.append("\n\\newpage\n")
    return clean_markdown("".join(parts)).strip() + "\n"

def build_docx_with_pandoc_to_path(md_text: str, out_path: Path, resource_dirs: Optional[List[Path]] = None) -> None:
    if not HAVE_PYPANDOC:
        raise RuntimeError("pypandoc not installed")
    if not ensure_pandoc_available():
        raise RuntimeError("Pandoc not available and auto-download failed")
    with tempfile.TemporaryDirectory() as td:
        md_file = Path(td) / "in.md"
        md_file.write_text(md_text, encoding="utf-8")
        extra_args = ["--standalone"]
        if resource_dirs:
            search_path = os.pathsep.join(str(Path(p).resolve()) for p in resource_dirs)
            extra_args.append(f"--resource-path={search_path}")
        pypandoc.convert_file(
            str(md_file),
            to="docx",
            format="gfm+tex_math_dollars+pipe_tables",
            outputfile=str(out_path),
            extra_args=extra_args,
        )

# ==== Table parsing for DOCX ===================================================
def parse_markdown_table(table_text: str) -> Tuple[List[List[str]], List[str]]:
    """Parse markdown table into rows, cells, and alignment info"""
    lines = [line.strip() for line in table_text.split('\n') if line.strip()]
    if not lines:
        return [], []
    
    rows = []
    alignments = []
    separator_line = None
    
    for line in lines:
        # Check if it's a separator line (like |:--:|:--:|)
        if re.match(r'^\|?\s*:?-+:?\s*\|', line):
            separator_line = line
            # Parse alignment from separator
            cells = [cell.strip() for cell in line.split('|')]
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            
            for cell in cells:
                if cell.startswith(':') and cell.endswith(':'):
                    alignments.append('center')
                elif cell.endswith(':'):
                    alignments.append('right')
                elif cell.startswith(':'):
                    alignments.append('left')
                else:
                    alignments.append('left')
            continue
        
        # Split by | and clean cells
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last if they exist (from leading/trailing |)
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        
        if cells:
            rows.append(cells)
    
    # If no alignment found, default to left for all columns
    if not alignments and rows:
        alignments = ['left'] * len(rows[0]) if rows else []
    
    return rows, alignments

def add_markdown_table(doc: Document, table_text: str):
    """Add a markdown table to DOCX document with better formatting"""
    rows, alignments = parse_markdown_table(table_text)
    if not rows:
        return
    
    # Find max columns across all rows
    max_cols = max(len(row) for row in rows) if rows else 1
    
    # Create table with correct dimensions
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill table cells with better formatting
    for i, row_data in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            cell_text = row_data[j] if j < len(row_data) else ""
            
            # Clean cell text (remove <br> tags, etc.)
            cell_text = cell_text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
            cell_text = re.sub(r'<[^>]+>', '', cell_text)  # Remove HTML tags
            
            # Set cell text
            cell.text = cell_text.strip()
            
            # Set alignment
            alignment = alignments[j] if j < len(alignments) else 'left'
            if alignment == 'center':
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Format first row as header (bold)
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Better cell formatting - set paragraph spacing
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.0

# ==== Basic DOCX fallback =====================================================
def add_text_block(doc: Document, text: str):
    """Add text to document, handling markdown formatting including tables"""
    # Clean the text first (but preserve table structure)
    # Don't clean tables yet - we'll handle them separately
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    
    # Split by double newlines for paragraphs
    chunks = text.split("\n\n")
    
    i = 0
    while i < len(chunks):
        chunk = chunks[i].strip()
        if not chunk:
            i += 1
            continue
        
        # Check if it's a heading
        if chunk.startswith('#'):
            # Count # to determine heading level
            level = len(chunk) - len(chunk.lstrip('#'))
            heading_text = chunk.lstrip('#').strip()
            if heading_text and level <= 6:
                doc.add_heading(heading_text, level=min(level, 6))
            i += 1
            continue
        
        # Check if it's a table (markdown table syntax)
        if '|' in chunk and chunk.count('|') >= 2:
            # Collect all table lines (including separator and continuation)
            table_lines = [chunk]
            j = i + 1
            while j < len(chunks):
                next_chunk = chunks[j].strip()
                # Check if it's a table line (has |) or separator line
                if next_chunk and ('|' in next_chunk or re.match(r'^\|?\s*:?-+:?\s*\|', next_chunk)):
                    table_lines.append(next_chunk)
                    j += 1
                else:
                    break
            
            # Combine table lines and create DOCX table
            table_text = '\n'.join(table_lines)
            add_markdown_table(doc, table_text)
            i = j
            continue
        
        # Check if it's a list (bulleted or numbered)
        if chunk.startswith('- ') or chunk.startswith('* ') or chunk.startswith('+ '):
            # Bulleted list
            list_items = [line[2:].strip() for line in chunk.split('\n') if line.strip().startswith(('- ', '* ', '+ '))]
            for item in list_items:
                if item:
                    p = doc.add_paragraph(item, style='List Bullet')
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Check if it's a numbered list
        numbered_match = re.match(r'^(\d+)[.)]\s+(.+)', chunk)
        if numbered_match:
            list_items = []
            for line in chunk.split('\n'):
                match = re.match(r'^(\d+)[.)]\s+(.+)', line.strip())
                if match:
                    list_items.append(match.group(2))
            for item in list_items:
                if item:
                    p = doc.add_paragraph(item, style='List Number')
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Regular paragraph - clean it before adding
        chunk_clean = clean_markdown_for_docx(chunk)
        if chunk_clean:
            p = doc.add_paragraph(chunk_clean)
            # Better paragraph spacing
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
        i += 1

def add_picture_fit_width(doc: Document, image_path: Path, max_width_in: float):
    doc.add_picture(str(image_path), width=Inches(max_width_in))

def build_docx_with_python_docx_to_path(pages_text: List[str],
                                        images_by_page: Dict[int, List[Path]],
                                        crops_by_page: Dict[int, List[Path]],
                                        out_path: Path,
                                        insert_page_breaks: bool,
                                        image_max_width_in: float) -> None:
    doc = Document()
    
    # Better default styles
    style = doc.styles['Normal'].font
    style.name = "Calibri"
    style.size = Pt(11)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Better paragraph style
    normal_style = doc.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15
    
    for i, txt in enumerate(pages_text, start=1):
        # Better heading formatting
        heading = doc.add_heading(f"Page {i}", level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
        if _strip(txt):
            add_text_block(doc, txt)
        
        for im in images_by_page.get(i, []):
            add_picture_fit_width(doc, im, image_max_width_in)
        
        for im in crops_by_page.get(i, []):
            add_picture_fit_width(doc, im, image_max_width_in)
        
        if insert_page_breaks and i < len(pages_text):
            doc.add_page_break()
    
    doc.save(out_path)

# ==== Regions JSON -> saved crops ============================================
def crops_from_regions(pdf_path: Path, regions_json_path: Path, dpi: int, assets_dir: Path) -> Dict[int, List[Path]]:
    with open(regions_json_path, "r", encoding="utf-8") as jf:
        cfg = json.load(jf)
    pages_cfg = cfg.get("pages", {})
    out: Dict[int, List[Path]] = {}
    for p_str, regs in pages_cfg.items():
        try:
            pnum = int(p_str)
        except:
            continue
        if not isinstance(regs, list):
            continue
        page_img, pts_wh = render_pdf_page_to_image(pdf_path, pnum, dpi=dpi)
        for k, r in enumerate(regs, start=1):
            coords     = r["coords"]
            coord_type = r.get("coord_type","norm")
            origin     = r.get("origin","top-left")
            padding    = int(r.get("padding", 8))
            label      = r.get("label", f"crop{k}")
            bbox_px = bbox_to_pixels(tuple(coords), page_img.width, page_img.height,
                                     pts_wh if coord_type=="pdf_points" else None,
                                     coord_type=coord_type, origin=origin)
            saved = crop_and_save(page_img, bbox_px, padding, assets_dir, f"p{pnum}_{k}_{label}")
            out.setdefault(pnum, []).append(saved)
    return out

# ==== Output path resolver (robust) ===========================================
def resolve_output_paths(pdf_path: Path, out_arg: Optional[str]) -> Tuple[Path, Path]:
    if out_arg:
        out = Path(out_arg)
        if out.suffix.lower() == ".docx":
            docx_path = out
            assets_dir = out.parent / (out.stem + "_assets")
        elif out.exists() and out.is_dir():
            docx_path = out / (pdf_path.stem + "_ocr.docx")
            assets_dir = out / (pdf_path.stem + "_assets")
        else:
            docx_path = Path(str(out) + "_ocr.docx")
            assets_dir = Path(str(out) + "_assets")
    else:
        docx_path = pdf_path.with_name(pdf_path.stem + "_ocr.docx")
        assets_dir = pdf_path.with_name(pdf_path.stem + "_assets")
    return docx_path, assets_dir

# ==== Helper to collect images (recursive + tolerant) =========================
def _iter_possible_images(p: Dict[str, Any]) -> List[Dict[str, Any]]:
    imgs: List[Dict[str, Any]] = []
    for key in ("images", "figures", "media", "inline_images"):
        val = p.get(key)
        if isinstance(val, list):
            imgs.extend([v for v in val if isinstance(v, dict)])

    def rec(node):
        if isinstance(node, dict):
            maybe_url = node.get("url") or node.get("src") or node.get("image")
            maybe_b64 = node.get("base64") or node.get("image_base64") or node.get("data") or node.get("imageData") or node.get("content") or node.get("b64")
            if isinstance(maybe_b64, str) and (len(maybe_b64) > 100 or maybe_b64.lower().startswith("data:")):
                imgs.append(node)
            elif isinstance(maybe_url, str) and (maybe_url.lower().startswith("data:") or maybe_url.lower().startswith("http")):
                imgs.append(node)
            for v in node.values():
                rec(v)
        elif isinstance(node, list):
            for v in node:
                rec(v)
    rec(p)

    seen = set()
    deduped = []
    for im in imgs:
        key = json.dumps({k: im.get(k) for k in ("id","url","src","image")}, sort_keys=True, default=str)
        if key not in seen:
            seen.add(key)
            deduped.append(im)
    return deduped

# ==== Main ====================================================================
def main():
    if not MISTRAL_OCR_ENDPOINT or not MISTRAL_API_KEY:
        die("Set MISTRAL_OCR_ENDPOINT and MISTRAL_API_KEY in your environment (.env).")

    ap = argparse.ArgumentParser(description="OCR PDF → Markdown (+ DOCX via Pandoc or python-docx, text-only, no images).")
    ap.add_argument("pdf", nargs="+", help="Path to the PDF (quote if it has spaces)")
    ap.add_argument("--title", default=None, help="Document title heading (markdown H1)")
    ap.add_argument("--out", default=None, help="Output markdown file path (default: PDF name with .md extension)")
    ap.add_argument(
        "--backend",
        choices=["pandoc", "python-docx"],
        default="python-docx",
        help="DOCX backend: 'pandoc' (better math) or 'python-docx' (better custom tables & formatting). Default: python-docx.",
    )
    ap.add_argument("--no-page-breaks", dest="no_page_breaks", action="store_true", help="No page breaks between pages")
    args = ap.parse_args()

    path = " ".join(args.pdf)
    pdf_path = Path(path)
    if not pdf_path.exists():
        die(f"File not found: {pdf_path}")

    pdf_bytes = pdf_path.read_bytes()
    print(f"[INFO] file={pdf_path.name} bytes={len(pdf_bytes)}")

    payload = {
        "model": MISTRAL_MODEL,
        "document": {"type": "document_url", "document_url": bytes_to_data_url("application/pdf", pdf_bytes)},
        "include_image_base64": False,  # Only text/markdown, no images from OCR
    }
    resp = post_ocr(payload)
    Path("ocr_response.json").write_text(json.dumps(resp, ensure_ascii=False, indent=2), encoding="utf-8")
    print("[INFO] wrote ocr_response.json")

    container = unwrap_container(resp)
    pages = container.get("pages")
    if not isinstance(pages, list) or not pages:
        top = ""
        for k in ("markdown","full_text","content","text","raw_text"):
            if isinstance(container.get(k), str) and container[k].strip():
                top = container[k]; break
        if not top.strip():
            die("No pages and no usable top-level text found in OCR response.")
        pages = [{"markdown": top}]

    # Determine markdown output path
    if args.out and args.out.endswith('.md'):
        md_path = Path(args.out)
    elif args.out:
        md_path = Path(str(args.out) + ".md")
    else:
        md_path = pdf_path.with_suffix(".md")
    print(f"[OUT] markdown={md_path}")
    print(f"[OUT] text-only mode: no images, no assets directory")

    pages_text: List[str] = []
    ocr_images_by_page: Dict[int, List[Path]] = {}

    print(f"[OK] pages={len(pages)}")
    for i, p in enumerate(pages, start=1):
        txt = extract_from_page(p if isinstance(p, dict) else {}) or ""
        if args.title and i == 1:
            txt = f"# {args.title}\n\n{txt}"
        pages_text.append(txt)

        # SKIP image collection - text only mode
        # Images are not collected or saved in this version
        ocr_images_by_page[i] = []

    print(f"[INFO] text-only mode: no images collected")

    # === SKIP REGION CROPS ===
    # No images/crops in text-only mode
    crops_by_page: Dict[int, List[Path]] = {}
    print(f"[INFO] text-only mode: no page crops generated")

    # Build markdown text
    md_text = build_markdown(
        pages_text,
        images_by_page=ocr_images_by_page,
        crops_by_page=crops_by_page,
        insert_page_breaks=not args.no_page_breaks,
        image_max_width_in=6.5,  # Default, not used since no images
    )

    # Save markdown file
    md_path.write_text(md_text, encoding="utf-8")
    print(f"[OK] wrote markdown: {md_path}")

    # Decide DOCX backend
    docx_path = md_path.with_suffix(".docx")
    backend = args.backend
    print(f"[INFO] DOCX backend requested: {backend}")

    if backend == "pandoc":
        # Option 1: Pandoc backend (better math, good tables if markdown is clean)
        try:
            print("[INFO] Using Pandoc backend for DOCX...")
            build_docx_with_pandoc_to_path(md_text, docx_path, resource_dirs=None)
            print(f"[OK] wrote docx (pandoc): {docx_path}")
        except Exception as e:
            # Fallback to python-docx so you always get a DOCX
            print(f"[WARN] Pandoc backend failed ({e}); falling back to python-docx.")
            build_docx_with_python_docx_to_path(
                pages_text,
                images_by_page=ocr_images_by_page,
                crops_by_page=crops_by_page,
                out_path=docx_path,
                insert_page_breaks=not args.no_page_breaks,
                image_max_width_in=6.5,
            )
            print(f"[OK] wrote docx (python-docx fallback): {docx_path}")
    else:
        # Option 2: python-docx backend (your current improved formatting & tables)
        print("[INFO] Using python-docx backend for DOCX...")
        build_docx_with_python_docx_to_path(
            pages_text,
            images_by_page=ocr_images_by_page,
            crops_by_page=crops_by_page,
            out_path=docx_path,
            insert_page_breaks=not args.no_page_breaks,
            image_max_width_in=6.5,
        )
        print(f"[OK] wrote docx (python-docx): {docx_path}")

if __name__ == "__main__":
    main()

